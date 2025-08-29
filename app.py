from fastapi import FastAPI, Request
from fastapi.responses import JSONResponse, FileResponse
from docx import Document
import uvicorn
import os
import datetime

app = FastAPI()

TEMPLATE_FILE = "Smart_Final_Board_Report_Template_WithForminatorIDs.docx"
OUTPUT_DIR = "generated_reports"

if not os.path.exists(OUTPUT_DIR):
    os.makedirs(OUTPUT_DIR)


def replace_placeholders(template_path, output_path, form_data):
    doc = Document(template_path)
    for para in doc.paragraphs:
        for key, value in form_data.items():
            if key in para.text:
                inline = para.runs
                for i in range(len(inline)):
                    if key in inline[i].text:
                        inline[i].text = inline[i].text.replace(key, str(value))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in form_data.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, str(value))

    doc.save(output_path)


@app.get("/")
def home():
    return {"status": "ok", "message": "Board Report Generator API is running",
            "endpoints": ["/generate-report/", "/download-report/{file_name}", "/debug/"]}


@app.get("/debug/")
async def debug(request: Request):
    body = await request.body()
    return {
        "headers": dict(request.headers),
        "query_params": dict(request.query_params),
        "raw_body": body.decode("utf-8") if body else "",
    }


@app.api_route("/generate-report/", methods=["GET", "POST"])
async def generate_report(request: Request):
    try:
        data = {}
        content_type = request.headers.get("content-type", "")

        if "application/json" in content_type:
            data = await request.json()
        elif "application/x-www-form-urlencoded" in content_type or "multipart/form-data" in content_type:
            form = await request.form()
            data = dict(form)
        else:
            raw_body = await request.body()
            if raw_body:
                data = {"raw_body": raw_body.decode("utf-8")}

        if not data:
            return JSONResponse({"status": "error", "message": "No form data received"}, status_code=400)

        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_file = os.path.join(OUTPUT_DIR, f"BoardReport_{timestamp}.docx")

        replace_placeholders(TEMPLATE_FILE, output_file, data)

        return {
            "status": "success",
            "message": "Report generated",
            "download_url": f"/download-report/{os.path.basename(output_file)}",
            "received_data": data
        }
    except Exception as e:
        return JSONResponse({"status": "error", "message": str(e)}, status_code=500)


@app.get("/download-report/{file_name}")
async def download_report(file_name: str):
    file_path = os.path.join(OUTPUT_DIR, file_name)
    if not os.path.exists(file_path):
        return JSONResponse({"status": "error", "message": "File not found"}, status_code=404)
    return FileResponse(path=file_path, filename=file_name, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=int(os.environ.get("PORT", 10000)))
